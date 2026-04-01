import logging

from django.contrib import messages
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse_lazy
from django.views.generic import CreateView, DetailView, ListView

from .forms import DocumentUploadForm
from .models import Document, OCRResult
from .services.ocr_engine import process_document

logger = logging.getLogger(__name__)


def index_redirect(request):
    """Redirect root to document list or upload."""
    return redirect('upload_document')


class DocumentUploadView(CreateView):
    model = Document
    form_class = DocumentUploadForm
    template_name = 'ocr/upload.html'
    success_url = reverse_lazy('document_list')

    def form_valid(self, form):
        # Save the original filename before saving the document
        form.instance.original_filename = form.cleaned_data['file'].name
        response = super().form_valid(form)
        
        # Process the document synchronously for now
        # In the future, this should be moved to a Celery task
        try:
            extracted_text = process_document(self.object)
            OCRResult.objects.create(
                document=self.object,
                extracted_text=extracted_text
            )
            messages.success(self.request, "Document processed successfully!")
        except Exception as e:
            logger.error(f"OCR failed during upload: {e}")
            messages.error(
                self.request,
                f"Failed to process document. Ensure it is a valid image or PDF."
            )
            
        return redirect('document_detail', pk=self.object.pk)


class DocumentListView(ListView):
    model = Document
    template_name = 'ocr/document_list.html'
    context_object_name = 'documents'
    paginate_by = 20


class DocumentDetailView(DetailView):
    model = Document
    template_name = 'ocr/document_detail.html'
    context_object_name = 'document'


from django.views.generic import DeleteView

class DocumentDeleteView(DeleteView):
    model = Document
    success_url = reverse_lazy('document_list')
    
    def delete(self, request, *args, **kwargs):
        messages.success(self.request, "Document deleted successfully.")
        return super().delete(request, *args, **kwargs)


def download_text(request, pk):
    """Download the extracted text in the requested format (.txt, .docx, .xlsx)."""
    document = get_object_or_404(Document, pk=pk)
    
    if not hasattr(document, 'result'):
        messages.error(request, "No text found for this document.")
        return redirect('document_detail', pk=pk)

    file_format = request.GET.get('format', 'txt').lower()
    content = document.result.extracted_text
    base_filename = f"{document.original_filename}_ocr"

    if file_format == 'docx':
        import io
        import os
        
        base_filename = document.original_filename.rsplit('.', 1)[0]
        
        if document.is_pdf:
            try:
                import tempfile
                from pdf2docx import Converter
                
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_out:
                    output_path = temp_out.name
                    
                cv = Converter(document.file.path)
                cv.convert(output_path, start=0, end=None)
                cv.close()

                with open(output_path, 'rb') as f:
                    docx_data = f.read()
                    
                os.remove(output_path)
                
                response = HttpResponse(
                    docx_data,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = f'attachment; filename="{base_filename}_converted.docx"'
                return response
            except Exception as e:
                logger.error(f"PDF to DOCX conversion failed: {e}")
                messages.error(request, "Rich PDF to Word conversion failed. Ensure the PDF is valid.")
                return redirect('document_detail', pk=pk)
                
        else:
            # Fallback for images: Create basic DOCX with extracted text
            from docx import Document as DocxDocument
            
            doc = DocxDocument()
            doc.add_heading(f'OCR Extraction: {document.original_filename}', 0)
            doc.add_paragraph(content)
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            response = HttpResponse(
                buffer.getvalue(), 
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{base_filename}_ocr.docx"'
            return response

    elif file_format == 'xlsx':
        import io
        import re
        from openpyxl import Workbook
        
        base_filename = document.original_filename.rsplit('.', 1)[0]
        
        # Strip HTML tags
        clean_content = re.sub(r'<[^>]+>', '', content)
        clean_content = clean_content.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')

        wb = Workbook()
        ws = wb.active
        ws.title = "OCR Results"
        
        # Write each line to a new row
        for row_idx, line in enumerate(clean_content.splitlines(), start=1):
            ws.cell(row=row_idx, column=1, value=line)
            
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(
            buffer.getvalue(), 
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{base_filename}_ocr.xlsx"'
        return response

    else:
        # Default to txt
        import re
        base_filename = document.original_filename.rsplit('.', 1)[0]
        # Strip HTML tags for plain text export
        clean_content = re.sub(r'<[^>]+>', '', content)
        # Decode common HTML entities
        clean_content = clean_content.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
        
        response = HttpResponse(clean_content, content_type='text/plain')
        response['Content-Disposition'] = f'attachment; filename="{base_filename}_ocr.txt"'
        return response


from django.http import JsonResponse
from django.views.decorators.http import require_POST
import json

@require_POST
def update_document_text(request, pk):
    """Update the extracted text for a document."""
    document = get_object_or_404(Document, pk=pk)
    
    if not hasattr(document, 'result'):
        return JsonResponse({'status': 'error', 'message': 'No OCR result found.'}, status=404)
        
    try:
        data = json.loads(request.body)
        new_text = data.get('text', '')
        
        document.result.extracted_text = new_text
        document.result.save()
        
        return JsonResponse({'status': 'success'})
    except Exception as e:
        logger.error(f"Failed to update text: {e}")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
